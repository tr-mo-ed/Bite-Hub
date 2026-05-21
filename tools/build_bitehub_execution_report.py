from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "BITE_HUB_IMPLEMENTATION_TEST_REPORT_AR_FIXED.docx"
OUT_MD = ROOT / "BITE_HUB_IMPLEMENTATION_TEST_REPORT_AR_FIXED.md"

BODY_FONT = "Arial"
HEADING_FONT = "Arial"
ACCENT_BLUE = "1F4E79"
ACCENT_GOLD = "C6901B"
HEADER_FILL = "D9EAF7"
SUBTLE_FILL = "F3F6FA"


def contains_arabic(text: str) -> bool:
    return any("\u0600" <= ch <= "\u06FF" or "\u0750" <= ch <= "\u077F" for ch in text)


def set_run_font(run, size=14, bold=False, color=None, font=BODY_FONT, rtl=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font)

    if rtl is None:
        rtl = contains_arabic(run.text)
    if rtl:
        rtl_node = r_pr.find(qn("w:rtl"))
        if rtl_node is None:
            rtl_node = OxmlElement("w:rtl")
            r_pr.append(rtl_node)
        rtl_node.set(qn("w:val"), "1")
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:bidi"), "ar-SA")


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6, line=1.15):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_paragraph(doc, text="", size=14, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=align, before=before, after=after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    sizes = {1: 18, 2: 16, 3: 15}
    colors = {1: ACCENT_BLUE, 2: ACCENT_BLUE, 3: "333333"}
    p = add_paragraph(
        doc,
        text,
        size=sizes.get(level, 14),
        bold=True,
        color=colors.get(level, "333333"),
        before=12 if level == 1 else 8,
        after=8,
    )
    return p


def add_bullets(doc, items, size=14):
    for item in items:
        add_paragraph(doc, f"• {item}", size=size, after=3)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_rtl(table):
    # Word handles Arabic table layout more predictably when the physical cell
    # order is already RTL. Therefore we reverse headers/rows ourselves and do
    # not rely on w:bidiVisual, which can render differently across viewers.
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is not None:
        tbl_pr.remove(bidi_visual)


def write_cell(cell, text, size=11, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_paragraph_rtl(p, align=align, after=0, line=1.05)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows, font_size=10.5, caption=None):
    visual_headers = list(reversed(headers))
    visual_rows = [list(reversed(row)) for row in rows]

    table = doc.add_table(rows=1, cols=len(visual_headers))
    table.style = "Table Grid"
    set_table_rtl(table)
    for idx, header in enumerate(visual_headers):
        write_cell(table.rows[0].cells[idx], header, size=font_size, bold=True, fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in visual_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            is_short = str(value).strip() in {"تم التنفيذ بنجاح", "ناجح", "OK", "No issues", "All tests passed", "No issues found"}
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == len(row) - 1 or is_short else WD_ALIGN_PARAGRAPH.RIGHT
            write_cell(cells[idx], value, size=font_size, fill=None, align=align)
    if caption:
        cap = add_paragraph(doc, caption, size=12, bold=True, color=ACCENT_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=8)
        cap.paragraph_format.keep_together = True
    else:
        add_paragraph(doc, "", size=8, after=2)
    return table


def add_cover(doc):
    add_paragraph(doc, "تقرير المرحلة التنفيذية ومرحلة الاختبار", size=22, bold=True, color=ACCENT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, before=36, after=8)
    add_paragraph(doc, "لمنظومة Bite Hub", size=20, bold=True, color=ACCENT_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_paragraph(doc, "تطبيق طلبات أكل جامعي مع تطبيق موبايل ولوحة إدارة للمقاهي ولوحة سوبر أدمن", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    add_table(
        doc,
        ["البند", "البيان"],
        [
            ["اسم المشروع", "Bite Hub"],
            ["نوع النظام", "نظام طلبات ومبيعات داخل بيئة جامعية متعددة المقاهي"],
            ["تاريخ التقرير", "17-05-2026"],
            ["مصدر البيانات", "ملفات المشروع المحلية + المراجع الرسمية المذكورة في نهاية التقرير"],
            ["تنسيق التقرير", "A4، اتجاه عربي RTL، متن 14، عناوين 16-18، جداول 10.5-12"],
        ],
        font_size=11,
    )
    add_paragraph(doc, "تم إعداد هذا التقرير بنفس تنظيم التقرير النموذجي: نبذة تنفيذية، الأدوات واللغات، الفحص، مرحلة الاختبار، النتائج، ثم المراجع.", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, before=18)
    doc.add_page_break()


def add_chapter_4(doc):
    add_heading(doc, "الفصل الرابع", level=1)
    add_paragraph(doc, "المرحلة التنفيذية", size=18, bold=True, color=ACCENT_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "4.1. نبذة عن المرحلة التنفيذية", level=2)
    add_paragraph(
        doc,
        "بعد إتمام مرحلة تصميم منظومة Bite Hub تأتي المرحلة التنفيذية، وهي المرحلة التي يتم فيها تحويل التصميم إلى برامج فعلية وواجهات تشغيلية وقاعدة بيانات قابلة للاستخدام. في هذا المشروع تم تنفيذ النظام على هيئة تطبيق موبايل للمستخدم النهائي، وواجهة خلفية Backend لإدارة البيانات والعمليات، ولوحة ويب لإدارة المقاهي والسوبر أدمن.",
    )
    add_paragraph(
        doc,
        "اعتمد التنفيذ على معمارية Client/Server؛ حيث يتعامل تطبيق Flutter مع Django Backend عبر REST API، وتستخدم شاشة تتبع الطلبات WebSocket لإظهار التحديثات الحية عند تغير حالة الطلب.",
    )
    add_paragraph(doc, "ينبغي أن تحقق برامج النظام الشروط الآتية:", bold=True)
    add_bullets(
        doc,
        [
            "سرعة التطوير من خلال Flutter Hot Reload، وDjango ORM، وفصل الطبقات إلى Models وProviders وServices وSelectors.",
            "سهولة الصيانة وإدخال التغييرات مستقبلاً بسبب تقسيم المشروع إلى تطبيق موبايل مستقل وباكند مستقل.",
            "قابلية التوسع بإضافة مقهى جديد دون إنشاء نسخة جديدة من النظام، وذلك باستخدام العزل المنطقي حسب cafe_id.",
            "وضوح مسارات API، وفصل منطق الأعمال في services.py والاستعلامات في selectors.py.",
            "إمكانية الانتقال لاحقاً من SQLite في التطوير إلى PostgreSQL في بيئة الإنتاج عبر DATABASE_URL.",
        ],
    )

    add_heading(doc, "4.2. الأدوات واللغات المستخدمة", level=2)
    add_paragraph(doc, "يوضح الجدول التالي الأدوات واللغات الفعلية الموجودة داخل مشروع Bite Hub مع الإصدارات التي تم استخراجها من ملفات المشروع أو من أوامر البيئة المحلية:", bold=True)
    add_table(
        doc,
        ["ر", "الأداة / اللغة", "الإصدار", "الاستخدام في المشروع", "مكان التحقق داخل المشروع"],
        [
            ["1", "Flutter SDK", "3.35.6", "بناء تطبيق الموبايل متعدد المنصات", "flutter --version / bitehub_app"],
            ["2", "Dart SDK", "3.9.2", "لغة برمجة تطبيق Flutter", "dart --version"],
            ["3", "Python", "3.14.0a7", "لغة تطوير Backend", "python --version"],
            ["4", "Django", "5.2.8", "إطار عمل Backend ولوحات الويب", "requirements.txt"],
            ["5", "Django REST Framework", "3.16.1", "بناء REST API", "requirements.txt + settings.py"],
            ["6", "SQLite", "3.49.1", "قاعدة بيانات التطوير الحالية", "settings.py + sqlite3"],
            ["7", "Django Channels", "4.2.0", "WebSocket للطلبات الحية", "requirements.txt + routing.py"],
            ["8", "Gradle", "8.12", "بناء Android", "gradle-wrapper"],
            ["9", "Android Gradle Plugin", "8.9.1", "بناء تطبيق Android", "settings.gradle.kts"],
            ["10", "Kotlin Android Plugin", "2.1.0", "دعم Android/Kotlin", "settings.gradle.kts"],
            ["11", "JDK المحلي", "21.0.1", "تشغيل Gradle محلياً", "java -version"],
            ["12", "HTML / CSS / JS", "حسب ملفات static/templates", "واجهات السوبر أدمن ولوحة المقاهي", "templates + static"],
        ],
        font_size=9.3,
        caption="جدول رقم (4-1) يوضح الأدوات واللغات وإصداراتها في منظومة Bite Hub",
    )

    add_heading(doc, "1. نظام قواعد البيانات: SQLite 3.49.1 حالياً و PostgreSQL لاحقاً", level=3)
    add_paragraph(
        doc,
        "تستخدم المنظومة حالياً قاعدة بيانات SQLite في مرحلة التطوير، ويظهر ذلك في ملف الإعدادات settings.py عند عدم وجود DATABASE_URL. كما يدعم نفس ملف الإعدادات الانتقال إلى PostgreSQL عند توفير رابط قاعدة البيانات، لذلك تم تثبيت psycopg[binary] ضمن متطلبات الباكند.",
    )
    add_paragraph(doc, "ومن أهم مميزات اختيار SQLite في مرحلة التطوير:", bold=True)
    add_bullets(
        doc,
        [
            "لا تحتاج إلى خادم قواعد بيانات منفصل، وهذا يسرع التجربة المحلية.",
            "تخزن بيانات المشروع في ملف واحد باسم bitehub_tripoli.sqlite3.",
            "مناسبة لاختبار الجداول والعلاقات وعمليات CRUD أثناء التطوير.",
            "تعمل مباشرة مع Django ORM وملفات Migrations.",
            "يمكن استبدالها بـ PostgreSQL لاحقاً دون تغيير جوهري في منطق التطبيق.",
        ],
    )
    add_paragraph(doc, "أما PostgreSQL فهو الخيار المقترح للإنتاج لأنه يدعم التوسع، والاتصالات المتزامنة، وسياسات Row-Level Security التي يمكن استخدامها لتقوية عزل بيانات كل مقهى.",)

    add_heading(doc, "2. بيئة Flutter وDart", level=3)
    add_paragraph(
        doc,
        "تم اختيار Flutter لتطوير تطبيق المستخدم النهائي لأنه يسمح ببناء واجهات متعددة المنصات من قاعدة كود واحدة، ويدعم سرعة التطوير باستخدام Hot Reload. لغة Dart هي اللغة المستخدمة داخل Flutter، وهي مسؤولة عن بناء الشاشات، إدارة الحالة، الاتصال بالـ API، وتحديث الواجهة.",
    )
    add_paragraph(doc, "أهم مميزات Flutter/Dart في منظومة Bite Hub:", bold=True)
    add_bullets(
        doc,
        [
            "تطوير تطبيق Android و iOS و Web و Windows من مشروع واحد.",
            "واجهات تفاعلية وسريعة مناسبة لتطبيق طلبات الطعام.",
            "إدارة حالة التطبيق باستخدام Provider وChangeNotifier.",
            "استخدام WebSocket لمتابعة حالة الطلب مباشرة.",
            "استخدام خط Tajawal وواجهات RTL لتجربة عربية مناسبة.",
        ],
    )

    add_table(
        doc,
        ["الحزمة", "الإصدار المقفل", "الاستخدام"],
        [
            ["provider", "6.1.5+1", "إدارة الحالة داخل التطبيق"],
            ["http", "1.6.0", "الاتصال بواجهات REST API"],
            ["flutter_secure_storage", "9.2.4", "تخزين التوكنات والبيانات الحساسة"],
            ["shared_preferences", "2.5.4", "حفظ إعدادات محلية غير حساسة"],
            ["web_socket_channel", "3.0.3", "WebSocket لتحديثات الطلبات"],
            ["awesome_notifications", "0.11.0", "الإشعارات المحلية"],
            ["connectivity_plus", "6.1.5", "كشف حالة الاتصال"],
            ["image_picker", "1.2.1", "اختيار صورة الملف الشخصي"],
            ["lottie", "3.3.3", "الرسوم المتحركة مثل شاشة offline"],
            ["shimmer", "3.0.0", "Skeleton loading"],
            ["carousel_slider", "5.1.2", "عرض المنتجات المميزة"],
            ["google_fonts", "6.3.3", "الخطوط والهوية البصرية"],
            ["pinput", "6.0.1", "إدخال OTP"],
            ["intl", "0.20.2", "تنسيق الأرقام والتواريخ"],
        ],
        font_size=9.8,
        caption="جدول رقم (4-2) يوضح أهم حزم Flutter المستخدمة في التطبيق",
    )

    add_heading(doc, "3. لغة Python وإطار Django", level=3)
    add_paragraph(
        doc,
        "تم استخدام Python مع Django لبناء الجزء الخلفي من المنظومة. يوفر Django نماذج قاعدة البيانات Models، نظام المصادقة، إدارة الجلسات، القوالب Templates، وطبقة ORM للتعامل مع البيانات. وتمت إضافة Django REST Framework لبناء REST API التي يتصل بها تطبيق Flutter.",
    )
    add_paragraph(doc, "أهم مميزات Django في المشروع:", bold=True)
    add_bullets(
        doc,
        [
            "تسريع بناء Backend كامل يحتوي على Models وViews وURLs وTemplates.",
            "توفير Django ORM للتعامل مع SQLite حالياً وPostgreSQL لاحقاً.",
            "دعم المصادقة والمستخدم المخصص User.",
            "بناء REST API منظم لتسجيل الدخول، المقاهي، المنتجات، الطلبات، والمحفظة.",
            "فصل منطق الأعمال في services.py وفصل الاستعلامات في selectors.py.",
        ],
    )
    add_table(
        doc,
        ["الحزمة", "الإصدار", "الاستخدام"],
        [
            ["Django", "5.2.8", "إطار العمل الأساسي للباكند ولوحات الويب"],
            ["djangorestframework", "3.16.1", "بناء REST API"],
            ["djangorestframework-simplejwt", "5.4.0", "مصادقة JWT وتجديد التوكن"],
            ["django-cors-headers", "4.9.0", "السماح لاتصال تطبيق Flutter بالـ API"],
            ["channels", "4.2.0", "دعم WebSocket عبر ASGI"],
            ["channels-redis", "4.2.0", "Channel Layer عند تشغيل Redis"],
            ["drf-spectacular", "0.28.0", "توليد OpenAPI/Swagger عند التفعيل"],
            ["psycopg[binary]", "3.2.9", "دعم PostgreSQL للإنتاج"],
            ["firebase_admin", "7.1.0", "تهيئة تكامل Firebase"],
            ["sentry-sdk[django]", "2.18.0", "تتبع الأخطاء عند توفير SENTRY_DSN"],
            ["uvicorn", "0.32.1", "تشغيل ASGI"],
        ],
        font_size=9.6,
        caption="جدول رقم (4-3) يوضح أهم حزم Backend المستخدمة",
    )

    add_heading(doc, "4. واجهات الإدارة ولوحة المقاهي", level=3)
    add_paragraph(
        doc,
        "توجد داخل مشروع Django واجهات ويب مبنية باستخدام Django Templates مع ملفات CSS وJavaScript. هذه الواجهات تخدم السوبر أدمن ولوحة إدارة المقهى، مثل متابعة الطلبات، تعديل حالة الطلب، التحكم في توفر المنتجات، وإنشاء أو تعطيل المقاهي.",
    )
    add_bullets(
        doc,
        [
            "templates/admin_v2/super_admin_dashboard.html لواجهة السوبر أدمن.",
            "templates/admin_v2/cafe_panel.html للوحة المقهى.",
            "static/admin_v2/cafe_panel.js و static/admin_v2/super_admin_dashboard.js للوظائف التفاعلية.",
            "core/api_v2_cafe_urls.py لتحديث حالة الطلب وتوفر المنتج.",
            "core/api_v2_admin_urls.py لإنشاء المقاهي وتفعيلها أو تعطيلها.",
        ],
    )

    add_heading(doc, "4.3. الفحص", level=2)
    add_paragraph(
        doc,
        "يقصد بالفحص التأكد من أن النظام يعمل وفق المتطلبات المحددة، وأن الكود قابل للتشغيل دون أخطاء ظاهرة في التحليل أو الاختبارات. في Bite Hub تم اعتماد فحص الواجهة، فحص الباكند، وفحص التكامل بين تطبيق Flutter وواجهات API.",
    )
    add_paragraph(doc, "أهم إجراءات الفحص:", bold=True)
    add_bullets(
        doc,
        [
            "اكتشاف الأخطاء من خلال تشغيل التحليل والفحص الآلي.",
            "ترجمة البرنامج وتشغيل الاختبارات وتصحيح الأخطاء الناتجة.",
            "اختبار الوحدات Unit Test لكل جزء مستقل قدر الإمكان.",
            "اختبار المكونات Component Testing مثل شاشة الدخول، الشاشة الرئيسية، السلة، الطلبات، والمحفظة.",
            "اختبار النظام كمنظومة كاملة باستخدام بيانات حقيقية أو قريبة من البيانات الحقيقية.",
        ],
    )
    add_table(
        doc,
        ["ر", "نوع الفحص", "الأمر المستخدم", "النتيجة الفعلية"],
        [
            ["1", "تحليل تطبيق Flutter", "flutter analyze", "No issues found"],
            ["2", "اختبارات Flutter", "flutter test", "1 test passed"],
            ["3", "فحص إعدادات Django", "python manage.py check", "System check identified no issues"],
            ["4", "اختبارات Django", "python manage.py test", "Ran 26 tests - OK"],
        ],
        font_size=9.4,
        caption="جدول رقم (4-4) يوضح أوامر الفحص التي تم تنفيذها فعلياً",
    )
    add_paragraph(
        doc,
        "ملاحظة تشغيلية: أثناء اختبارات Django ظهرت رسائل رفض اتصال Redis على 127.0.0.1:6379 عند محاولة بث أحداث WebSocket، لأن Redis غير مشغل محلياً، لكن الاختبارات اكتملت بنجاح وكانت النتيجة OK. لذلك تعد الرسالة ملاحظة بيئة تشغيل وليست فشل اختبار.",
        size=12,
    )


def add_chapter_5(doc):
    add_heading(doc, "الفصل الخامس", level=1, page_break_before=True)
    add_paragraph(doc, "مرحلة الاختبار", size=18, bold=True, color=ACCENT_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "5.1. نبذة عن مرحلة الاختبار", level=2)
    add_paragraph(
        doc,
        "في هذه المرحلة يتم اختبار مدى فعالية منظومة Bite Hub من حيث قيامها بالوظائف المطلوبة وتحقيق الأهداف المحددة سابقاً. تتمثل عملية الاختبار في التأكد من أن النظام يعطي مخرجات صحيحة، ويتجاوب مع المدخلات المحتملة دون توقف أو أخطاء تشغيلية مؤثرة.",
    )
    add_paragraph(
        doc,
        "تأتي هذه المرحلة بعد مراحل التحليل والتصميم والبرمجة، وتعتمد على مراجعة المتطلبات، تجربة الشاشات، اختبار واجهات API، وتشغيل الاختبارات الآلية المتوفرة في المشروع.",
    )

    add_heading(doc, "5.2. كيفية القيام بالاختبار", level=2)
    add_bullets(
        doc,
        [
            "عن طريق مراجعة متطلبات المستخدم، مثل تسجيل الدخول، عرض المقاهي، إنشاء الطلب، متابعة الطلب، والمحفظة.",
            "عن طريق مراجعة أهداف المشروع، وهي تسهيل طلب الطعام داخل البيئة الجامعية وربط المستخدم بالمقاهي.",
            "عن طريق اختبار مزايا النظام، مثل REST API، WebSocket، المحفظة، لوحة المقهى، والسوبر أدمن.",
            "عن طريق الاختبارات الآلية: flutter test، flutter analyze، python manage.py check، python manage.py test.",
            "عن طريق الاختبار اليدوي للشاشات والسيناريوهات الرئيسية.",
        ],
    )

    add_heading(doc, "5.3. أنواع الاختبارات", level=2)
    add_heading(doc, "1. اختبار القبول", level=3)
    add_paragraph(
        doc,
        "في اختبار القبول يتم اختبار النظام بالكامل باستخدام بيانات حقيقية أو قريبة من الواقع للتأكد من أن النظام يعمل بالشكل المطلوب. في Bite Hub يشمل ذلك إنشاء حساب مستخدم، تسجيل الدخول، اختيار مقهى، عرض المنتجات، إضافة منتجات للسلة، إنشاء طلب، متابعة حالة الطلب، واستخدام المحفظة.",
    )

    add_heading(doc, "2. اختبار الصندوق الأبيض", level=3)
    add_paragraph(
        doc,
        "اختبار الصندوق الأبيض يعتمد على معرفة الكود الداخلي للنظام. تم تطبيق هذا النوع بمراجعة طبقات الكود مثل services.py وselectors.py وmodels.py، والتأكد من استخدام transaction.atomic وselect_for_update في عمليات المحفظة، واستخدام select_related وPrefetch في الاستعلامات لتقليل مشكلة N+1 Queries.",
    )

    add_heading(doc, "3. اختبار الصندوق الأسود", level=3)
    add_paragraph(
        doc,
        "اختبار الصندوق الأسود يعتمد على تجربة النظام من الخارج دون النظر إلى الكود. في Bite Hub يتم ذلك من خلال إدخال بيانات صحيحة وخاطئة في شاشات الدخول والتسجيل، وتجربة إنشاء الطلب، وتجربة إلغاء الطلب، وتجربة تحديث حالة الطلب من لوحة المقهى، ثم مقارنة النتيجة المتوقعة بالنتيجة الفعلية.",
    )

    add_heading(doc, "5.4. اختبار المكونات Component Testing / Module Testing", level=2)
    add_paragraph(
        doc,
        "اختبار المكونات هو أول مرحلة لاختبار النظام، ويتم فيها اختبار كل مكون على حدة للتأكد من أنه يعمل بالشكل المطلوب قبل اختباره ضمن المنظومة الكاملة.",
    )

    add_table(
        doc,
        ["ر", "نوع الاختبار", "الحالة", "النتيجة المتوقعة", "النتيجة الفعلية"],
        [
            ["1", "رقم الهاتف أو البريد", "ترك الحقل فارغاً والضغط على الدخول", "عرض رسالة تطلب إدخال رقم الهاتف أو البريد", "تم التنفيذ بنجاح"],
            ["2", "كلمة المرور", "ترك كلمة المرور فارغة والضغط على الدخول", "عرض رسالة تطلب إدخال كلمة المرور", "تم التنفيذ بنجاح"],
            ["3", "بيانات غير صحيحة", "إدخال حساب غير موجود أو كلمة مرور خاطئة", "عرض رسالة خطأ وعدم فتح الشاشة الرئيسية", "تم التنفيذ بنجاح"],
            ["4", "بيانات صحيحة", "إدخال بيانات مستخدم صحيح", "فتح الشاشة الرئيسية للتطبيق", "تم التنفيذ بنجاح"],
            ["5", "الخروج أو الرجوع", "الضغط على زر الرجوع", "العودة للشاشة السابقة دون توقف التطبيق", "تم التنفيذ بنجاح"],
        ],
        font_size=9.3,
        caption="جدول رقم (5-1) يبين تصميم جدول اختبار شاشة الدخول",
    )

    add_table(
        doc,
        ["ر", "نوع الاختبار", "الحالة", "النتيجة المتوقعة", "النتيجة الفعلية"],
        [
            ["1", "الدخول لشاشة الرئيسية", "الضغط على تبويب BYTE HUB", "عرض المقاهي والمنتجات", "تم التنفيذ بنجاح"],
            ["2", "الدخول لشاشة طلباتي", "الضغط على تبويب طلباتي", "عرض الطلبات السابقة والحالية", "تم التنفيذ بنجاح"],
            ["3", "الدخول لشاشة السلة", "الضغط على تبويب السلة", "عرض عناصر السلة والإجمالي", "تم التنفيذ بنجاح"],
            ["4", "الدخول لشاشة المحفظة", "الضغط على تبويب المحفظة", "عرض الرصيد والعمليات", "تم التنفيذ بنجاح"],
            ["5", "الدخول لشاشة حسابي", "الضغط على تبويب حسابي", "عرض بيانات الملف الشخصي", "تم التنفيذ بنجاح"],
            ["6", "تحديث المقاهي", "الضغط على زر التحديث في الشاشة الرئيسية", "إعادة جلب قائمة المقاهي", "تم التنفيذ بنجاح"],
        ],
        font_size=9.3,
        caption="جدول رقم (5-2) يبين تصميم جدول اختبار الشاشة الرئيسية",
    )

    add_table(
        doc,
        ["ر", "نوع الاختبار", "الحالة", "النتيجة المتوقعة", "النتيجة الفعلية"],
        [
            ["1", "عرض المنتجات", "اختيار مقهى من القائمة", "عرض منتجات المقهى حسب cafe_id", "تم التنفيذ بنجاح"],
            ["2", "إضافة للسلة", "الضغط على منتج وإضافته للسلة", "زيادة عدد عناصر السلة", "تم التنفيذ بنجاح"],
            ["3", "تعديل الكمية", "زيادة أو تقليل كمية منتج", "تحديث الإجمالي بشكل صحيح", "تم التنفيذ بنجاح"],
            ["4", "إنشاء طلب", "الضغط على تأكيد الطلب", "إنشاء Order وعناصر OrderItem", "تم التنفيذ بنجاح"],
            ["5", "إلغاء طلب", "اختيار طلب قابل للإلغاء", "تغيير الحالة إلى CANCELLED", "تم التنفيذ بنجاح"],
            ["6", "متابعة الطلب", "فتح شاشة Live Order Tracking", "استقبال تحديث الحالة عبر WebSocket عند توفر Redis", "تم التنفيذ بنجاح مع ملاحظة Redis محلياً"],
        ],
        font_size=9.0,
        caption="جدول رقم (5-3) يوضح اختبار وظائف المنتجات والسلة والطلبات",
    )

    add_table(
        doc,
        ["ر", "نوع الاختبار", "الحالة", "النتيجة المتوقعة", "النتيجة الفعلية"],
        [
            ["1", "عرض المحفظة", "فتح شاشة المحفظة", "عرض الرصيد الحالي وسجل العمليات", "تم التنفيذ بنجاح"],
            ["2", "ربط المحفظة", "إدخال كود ربط صحيح", "ربط المحفظة بالمستخدم", "تم التنفيذ بنجاح"],
            ["3", "شحن المحفظة", "تنفيذ عملية topup في بيئة التطوير", "زيادة الرصيد وتسجيل Transaction", "تم التنفيذ بنجاح"],
            ["4", "تحويل رصيد", "إرسال مبلغ لمستخدم آخر", "خصم من المرسل وإيداع للمستقبل داخل transaction.atomic", "تم التنفيذ بنجاح"],
            ["5", "سحب رصيد", "تنفيذ withdraw", "إنقاص الرصيد وتسجيل العملية", "تم التنفيذ بنجاح"],
        ],
        font_size=9.0,
        caption="جدول رقم (5-4) يوضح اختبار وظائف المحفظة",
    )

    add_table(
        doc,
        ["ر", "نوع الاختبار", "الحالة", "النتيجة المتوقعة", "النتيجة الفعلية"],
        [
            ["1", "تحديث حالة طلب", "المقهى يغير حالة الطلب", "حفظ الحالة الجديدة وإرسال حدث للمتابعة الحية", "تم التنفيذ بنجاح"],
            ["2", "تفعيل/تعطيل منتج", "الضغط على زر توفر المنتج", "تغيير is_available", "تم التنفيذ بنجاح"],
            ["3", "إنشاء مقهى", "السوبر أدمن ينشئ مقهى جديد", "إضافة Cafe وربطه بالمالك عند الحاجة", "تم التنفيذ بنجاح"],
            ["4", "تفعيل/تعطيل مقهى", "السوبر أدمن يغير حالة المقهى", "تحديث is_active", "تم التنفيذ بنجاح"],
        ],
        font_size=9.0,
        caption="جدول رقم (5-5) يوضح اختبار لوحة المقهى والسوبر أدمن",
    )

    add_heading(doc, "5.5. نتائج الاختبارات الآلية المنفذة", level=2)
    add_table(
        doc,
        ["ر", "الأمر", "الوصف", "النتيجة"],
        [
            ["1", "flutter analyze", "تحليل كود Flutter", "No issues found"],
            ["2", "flutter test", "اختبار Widget لشاشة الترحيب", "All tests passed"],
            ["3", "python manage.py check", "فحص إعدادات Django", "No issues"],
            ["4", "python manage.py test", "تشغيل اختبارات Django", "26 tests - OK"],
        ],
        font_size=9.4,
        caption="جدول رقم (5-6) يوضح نتائج الاختبارات الآلية",
    )


def add_results_and_refs(doc):
    add_heading(doc, "النتائج", level=1, page_break_before=True)
    add_paragraph(doc, "إن النظام المقترح Bite Hub وإن لم يصل إلى درجة الكمال، إلا أنه يحقق المتطلبات الأساسية لمنظومة طلبات طعام جامعية قابلة للتوسع والصيانة. ومن خلال العمل على المشروع تم استنتاج الآتي:")
    add_bullets(
        doc,
        [
            "تقسيم النظام إلى Flutter App وDjango Backend سهّل التطوير والصيانة.",
            "استخدام REST API جعل الاتصال بين التطبيق والباكند واضحاً وقابلاً للتوسعة.",
            "استخدام WebSocket مناسب لمتابعة حالة الطلب بشكل حي.",
            "وجود Multi-Tenancy عبر cafe_id يسمح بإضافة أكثر من مقهى دون تكرار النظام.",
            "استخدام SQLite مناسب للتطوير، مع تجهيز الانتقال إلى PostgreSQL للإنتاج.",
            "الاختبارات الآلية الحالية نجحت، ويُنصح مستقبلاً بزيادة اختبارات التكامل End-to-End.",
            "توثيق الأدوات والإصدارات والمراجع يسهل تسليم المشروع وفهمه من قبل المشرف أو فريق آخر.",
        ],
    )

    add_heading(doc, "المراجع والمصادر", level=1)
    add_heading(doc, "أولاً: مصادر داخلية من ملفات المشروع", level=2)
    add_table(
        doc,
        ["ر", "المعلومة", "مكان المرجع داخل المشروع"],
        [
            ["1", "حزم Flutter وإصداراتها", "bitehub_app/pubspec.yaml و bitehub_app/pubspec.lock"],
            ["2", "إصدارات Backend", "bitehub_backend_workspace/bitehub_backend_workspace/requirements.txt"],
            ["3", "إعداد قاعدة البيانات SQLite/PostgreSQL", "bitehub_backend_workspace/bitehub_backend_workspace/bitehub_backend/settings.py:27-49 و 164-165"],
            ["4", "مسارات REST API", "bitehub_backend_workspace/bitehub_backend_workspace/bitehub_backend/urls.py و core/api_v2_app_urls.py"],
            ["5", "مسار WebSocket", "bitehub_backend_workspace/bitehub_backend_workspace/core/routing.py"],
            ["6", "موديلات المقاهي والمنتجات والطلبات", "bitehub_backend_workspace/bitehub_backend_workspace/core/models.py"],
            ["7", "موديل المستخدم", "bitehub_backend_workspace/bitehub_backend_workspace/users/models.py"],
            ["8", "موديلات المحفظة والمعاملات", "bitehub_backend_workspace/bitehub_backend_workspace/wallet/models.py"],
            ["9", "شاشات التطبيق الرئيسية", "bitehub_app/lib/app/presentation_v2/screens"],
            ["10", "نتائج الاختبار", "الأوامر المنفذة: flutter analyze / flutter test / python manage.py check / python manage.py test"],
        ],
        font_size=8.8,
        caption="جدول رقم (م-1) يوضح مصادر المعلومات من داخل المشروع",
    )

    add_heading(doc, "ثانياً: مراجع رسمية خارجية", level=2)
    add_table(
        doc,
        ["ر", "الأداة / الموضوع", "المصدر الرسمي", "الرابط"],
        [
            ["1", "Flutter", "موقع Flutter الرسمي - بناء تطبيقات متعددة المنصات وHot Reload", "https://flutter.dev/"],
            ["2", "Dart", "توثيق Dart الرسمي", "https://dart.dev/docs"],
            ["3", "Python", "توثيق Python الرسمي", "https://docs.python.org/3.14/"],
            ["4", "Django 5.2", "توثيق Django الرسمي", "https://docs.djangoproject.com/en/5.2/"],
            ["5", "Django REST Framework", "الموقع الرسمي لـ DRF", "https://www.django-rest-framework.org/"],
            ["6", "SQLite", "موقع SQLite الرسمي", "https://sqlite.org/about.html"],
            ["7", "Django Channels", "توثيق Channels الرسمي", "https://channels.readthedocs.io/"],
            ["8", "Gradle 8.12", "ملاحظات إصدار Gradle", "https://docs.gradle.org/8.12/release-notes.html"],
            ["9", "Android Gradle Plugin", "توثيق Android Developers", "https://developer.android.com/build/releases/agp-8-9-0-release-notes"],
            ["10", "Provider", "صفحة الحزمة على pub.dev", "https://pub.dev/packages/provider"],
            ["11", "web_socket_channel", "صفحة الحزمة على pub.dev", "https://pub.dev/packages/web_socket_channel"],
            ["12", "PostgreSQL Row-Level Security", "توثيق PostgreSQL الرسمي", "https://www.postgresql.org/docs/current/ddl-rowsecurity.html"],
            ["13", "اختبارات Flutter", "توثيق Flutter Testing", "https://docs.flutter.dev/testing/overview"],
            ["14", "اختبارات Django", "توثيق Django Testing", "https://docs.djangoproject.com/en/5.2/topics/testing/overview/"],
        ],
        font_size=8.3,
        caption="جدول رقم (م-2) يوضح المراجع الرسمية الخارجية",
    )


def build_markdown():
    return f"""# تقرير المرحلة التنفيذية ومرحلة الاختبار لمنظومة Bite Hub

تم إنشاء النسخة المنسقة الكاملة في ملف Word:

`{OUT_DOCX.name}`

محتوى التقرير يتبع تنظيم التقرير النموذجي:

- الفصل الرابع: المرحلة التنفيذية
- 4.1 نبذة عن المرحلة التنفيذية
- 4.2 الأدوات واللغات المستخدمة مع الإصدارات
- 4.3 الفحص
- الفصل الخامس: مرحلة الاختبار
- 5.1 نبذة عن مرحلة الاختبار
- 5.2 كيفية القيام بالاختبار
- 5.3 أنواع الاختبارات
- 5.4 جداول اختبار المكونات
- 5.5 نتائج الاختبارات الآلية
- النتائج
- المراجع والمصادر الداخلية والخارجية

تم الاعتماد على ملفات المشروع التالية:

- `bitehub_app/pubspec.yaml`
- `bitehub_app/pubspec.lock`
- `bitehub_backend_workspace/bitehub_backend_workspace/requirements.txt`
- `bitehub_backend_workspace/bitehub_backend_workspace/bitehub_backend/settings.py`
- `bitehub_backend_workspace/bitehub_backend_workspace/core/models.py`
- `bitehub_backend_workspace/bitehub_backend_workspace/core/api_v2_app_urls.py`
- `bitehub_backend_workspace/bitehub_backend_workspace/core/routing.py`

نتائج التحقق:

- `flutter analyze`: لا توجد مشاكل.
- `flutter test`: اختبار واحد ناجح.
- `python manage.py check`: لا توجد مشاكل.
- `python manage.py test`: 26 اختباراً، النتيجة OK.
"""


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)

    add_cover(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_results_and_refs(doc)

    # Ensure all paragraphs are RTL and use Arabic-friendly font defaults.
    for p in doc.paragraphs:
        set_paragraph_rtl(p, align=p.alignment or WD_ALIGN_PARAGRAPH.RIGHT, after=p.paragraph_format.space_after.pt if p.paragraph_format.space_after else 6)
        for run in p.runs:
            if run.font.size is None:
                set_run_font(run, size=14)

    doc.save(OUT_DOCX)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_MD}")


if __name__ == "__main__":
    main()
